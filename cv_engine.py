import os
import sys
import json
import re
import warnings
import copy
import platform
import subprocess
import time
import shutil
import datetime
def _extract_contacts_plus(raw_text: str) -> dict:
    """Best-effort extraction of email/phone/website/linkedin from raw text."""
    out = {"email":"", "phone":"", "website":"", "linkedin":""}
    if not isinstance(raw_text, str) or not raw_text.strip():
        return out
    t = raw_text

    m = re.search(r"([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})", t, re.I)
    if m: out["email"] = m.group(1)

    for m in re.finditer(r"(\+?\d[\d\s().-]{7,}\d)", t):
        cand = m.group(1)
        digits = re.sub(r"\D", "", cand)
        if len(digits) >= 9:
            out["phone"] = cand.strip()
            break

    m = re.search(r"(https?://(?:[a-z]{2,3}\.)?linkedin\.com/[\w\-./?=&%#]+)", t, re.I)
    if m: out["linkedin"] = m.group(1)

    m = re.search(r"(https?://(?![^\s]*linkedin\.com)[^\s)]+)", t, re.I)
    if m:
        out["website"] = m.group(1)
    else:
        m = re.search(r"\b(www\.[^\s)]+)", t, re.I)
        if m:
            out["website"] = m.group(1)

    return out

def _extract_location_line(raw_text: str) -> str:
    """Try to find a short location-like line near the top."""
    if not isinstance(raw_text, str) or not raw_text.strip():
        return ""
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    for ln in lines[:25]:
        if len(ln) > 60:
            continue
        low = ln.lower()
        if any(k in low for k in ("summary","experience","education","skills","top skills","certification","linkedin","http","www","@", "project:")):
            continue
        if "," in ln:
            return ln
    return ""



# --- 1. BASIC SETTINGS ---


def _trim_strings_deep(value):
    """Recursively trim whitespace and XML-escape strings for DOCX rendering.
    Note: &, < and > must be escaped because docxtpl does not autoescape Jinja2
    variables, so raw XML-special chars break DOCX XML (truncate content).
    """
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    if isinstance(value, list):
        return [_trim_strings_deep(v) for v in value]
    if isinstance(value, dict):
        return {_trim_strings_deep(k) if isinstance(k, str) else k: _trim_strings_deep(v) for k, v in value.items()}
    return value

def _normalize_optional_section_title(title):
    title = str(title or "").strip()
    if not title:
        return ""
    if "_" in title:
        title = title.replace("_", " ")
    if title.isupper() and len(title) > 4:
        title = title.title()
    return title

warnings.filterwarnings("ignore")
os.environ["GRPC_VERBOSITY"] = "ERROR"
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "3"

try:
    from google import genai as google_genai_sdk
    from google.genai import types as google_genai_types
except Exception:
    google_genai_sdk = None
    google_genai_types = None

def _make_google_genai_client(api_key):
    if not google_genai_sdk:
        raise RuntimeError('google-genai SDK is not installed or failed to import')
    return google_genai_sdk.Client(api_key=api_key)

def _retry_google_call(callable_fn):
    max_retries = 3
    delay = 5
    for attempt in range(max_retries):
        try:
            return callable_fn()
        except Exception as e:
            err_str = str(e)
            if "429" in err_str or "Resource exhausted" in err_str or "Quota" in err_str:
                if attempt < max_retries - 1:
                    print(f"⚠️ API 429 Limit hit. Sleeping for {delay} seconds... (Attempt {attempt+1}/{max_retries})")
                    time.sleep(delay)
                else:
                    raise e
            else:
                raise e

def _generate_text_content_with_retry(api_key, contents):
    client = _make_google_genai_client(api_key)
    return _retry_google_call(lambda: client.models.generate_content(model=MODEL_NAME, contents=contents))

def _get_gemini_file_state_name(file_obj):
    state = getattr(file_obj, 'state', None)
    if state is None:
        return ''
    name = getattr(state, 'name', None)
    if name:
        return str(name)
    return str(state)


def _upload_gemini_file_and_wait(api_key, file_path, mime_type, task_state=None):
    client = _make_google_genai_client(api_key)
    if google_genai_types is None:
        raise RuntimeError('google-genai types are not available')

    import unicodedata
    _safe_name = unicodedata.normalize('NFKD', os.path.basename(file_path)).encode('ascii', 'ignore').decode('ascii') or 'file'
    upload_config = google_genai_types.UploadFileConfig(mime_type=mime_type, display_name=_safe_name) if mime_type else None

    with open(file_path, 'rb') as _fh:
        upload_kwargs = {'file': _fh}
        if upload_config:
            upload_kwargs['config'] = upload_config
        sample = client.files.upload(**upload_kwargs)

    upload_wait = 0
    while _get_gemini_file_state_name(sample) == 'PROCESSING':
        if task_state and task_state.get('cancel'):
            return None
        time.sleep(1)
        upload_wait += 1
        if upload_wait > FILE_UPLOAD_TIMEOUT_SEC:
            raise TimeoutError(f"File upload timed out after 5 min: {os.path.basename(file_path)}")
        sample = client.files.get(name=sample.name)
    return sample


def _generate_file_content_with_retry(api_key, prompt_text, sample):
    client = _make_google_genai_client(api_key)
    return _retry_google_call(lambda: client.models.generate_content(model=MODEL_NAME, contents=[sample, prompt_text]))

def _extract_token_usage(response):
    usage = getattr(response, 'usage_metadata', None)
    if usage is None:
        return 0, 0
    in_tok = getattr(usage, 'prompt_token_count', 0) or 0
    out_tok = getattr(usage, 'candidates_token_count', 0) or getattr(usage, 'response_token_count', 0) or 0
    return in_tok, out_tok

try:
    from docxtpl import DocxTemplate
except ImportError:
    print("CRITICAL: docxtpl not installed. Please run: pip install docxtpl")

# ==========================================
# 🛑 DOCX "GHOST FOLDER" FIX
# ==========================================
def fix_docx_path_bug():
    if getattr(sys, 'frozen', False):
        try:
            base_dir = sys._MEIPASS
            parts_dir = os.path.join(base_dir, 'docx', 'parts')
            os.makedirs(parts_dir, exist_ok=True)
            templates_dir = os.path.join(base_dir, 'docx', 'templates')
            os.makedirs(templates_dir, exist_ok=True)
            header = os.path.join(templates_dir, 'default-header.xml')
            if not os.path.exists(header):
                with open(header, 'wb') as f:
                    f.write(b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n<w:hdr xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"></w:hdr>")
            footer = os.path.join(templates_dir, 'default-footer.xml')
            if not os.path.exists(footer):
                with open(footer, 'wb') as f:
                    f.write(b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n<w:ftr xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"></w:ftr>")
        except Exception: pass  # Non-critical: frozen bundle path fix

fix_docx_path_bug()
# ==========================================

MODEL_NAME = 'gemini-2.5-flash'
APP_VERSION = "03.49"

# 💸 GEMINI 2.0 FLASH PRICING (Per 1 Million Tokens)
PRICE_1M_IN = 0.15
PRICE_1M_OUT = 0.60
FILE_UPLOAD_TIMEOUT_SEC = 300
DOCX_SAVE_MAX_RETRIES = 20

def get_resource_path(relative_path):
    if getattr(sys, 'frozen', False):
        try:
            base_path = sys._MEIPASS
            if os.path.exists(os.path.join(base_path, relative_path)):
                return os.path.join(base_path, relative_path)
        except Exception: pass
        app_resources_path = os.path.join(os.path.dirname(sys.executable), '..', 'Resources')
        if os.path.exists(os.path.join(app_resources_path, relative_path)):
            return os.path.join(app_resources_path, relative_path)
    return os.path.join(os.path.abspath(os.path.dirname(__file__)), relative_path)

USER_HOME = os.path.expanduser("~")
DEFAULT_WORKSPACE = os.path.join(USER_HOME, "Documents", "Quantori_CV_Workplace")
SCRIPT_NAME = "QuantoriCV" if getattr(sys, 'frozen', False) else "CV Manager"
SETTINGS_FILE = os.path.join(USER_HOME, '.quantoricv_settings.json')
MASTER_PROMPTS_FILE = os.path.join(USER_HOME, '.master_prompts.json')

# ==========================================
# 2. PROMPTS & SCHEMAS (PROMPT LAB FOUNDATION)
# ==========================================

# 🛡️ PROTECTED SCHEMA: Never exposed to the user for editing
CV_JSON_SCHEMA = """{
  "basics": {
    "name": "String", "current_title": "String", "objective": "String",
    "contacts": { "email": "String", "phone": "String", "location": "String" },
    "links": ["String"]
  },
  "summary": { "bullet_points": ["String"] },
  "skills": { "Category": ["Skill 1"] },
  "experience": [{
      "category": "String", "company_name": "String", "role": "String",
      "dates": { "start": "String — date or duration-only e.g. '6 years 5 months'", "end": "String — date, 'Present', or '' if only duration" },
      "location": "String", "project_description": "String",
      "highlights": ["String"], "environment": ["String"]
  }],
  "education": [{"institution": "String", "degree": "String", "year": "String", "details": "String"}],
  "certifications": ["String"],
  "languages": [{"language": "String", "proficiency": "String", "level": "String", "details": "String"}],
  "other_sections": [{"title": "String", "items": ["String"]}]
}"""

# 🧠 MUTABLE PROMPTS: Available for editing in the interface
DEFAULT_PROMPTS = {
    "prompt_master_inst": """You are a STRICT Lossless CV Extractor and Translator. Your ONLY job is to digitize this CV into JSON matching the required schema with minimal data loss.

**CRITICAL RULES FOR LOSSLESS EXTRACTION (STRICTLY ENFORCED):**
1. **US ENGLISH ONLY:** All human-readable output in the JSON must be translated into professional US English. Do not leave Russian or other non-English prose in the output.
2. **NO INVENTED FACTS / NO DATA LOSS:** Extract only facts supported by the CV, but do not lose explicit information. Preserve meaningful technical terms, methods, tools, technologies, responsibilities, achievements, project details, and bullet points. **HIGHLIGHTS INTEGRITY:** Each item in a role's `highlights` array MUST correspond to an actual bullet point, responsibility, or achievement explicitly written in the source CV for that role. NEVER generate, rephrase into new meaning, or fabricate highlights. If a role has no bullets or descriptions in the source, return an empty `highlights: []`.
3. **DEEP SCAN THE ENTIRE CV:** Extract from the whole document, including header, summary, skills blocks, Top Skills, experience bullets, project descriptions, certifications, languages, links, and side sections.
4. **SKILLS & ENVIRONMENT:** Extract explicit skills, tools, technologies, frameworks, platforms, databases, cloud/services, and domain systems from all relevant sections. Put them into `skills` or role `environment` as appropriate. Do not create noisy or redundant generic skill categories.
5. **DATES & CURRENT STATUS:** Extract all explicit dates from ALL sections including education, preserving the highest precision supported by the source (`April 2025`, `2018`, `Present`). Education years (e.g. `2012 - 2016`) MUST go into `education[].year`. If only a duration is available with no start/end dates (e.g. `6 years 5 months`), place it in `dates.start` and leave `dates.end` empty. Never invent dates. Never assign future dates to finished past roles.
6. **CONTACTS, LINKS, LOCATIONS:** Extract all explicit phone numbers, emails, LinkedIn, GitHub, portfolio, websites, WhatsApp, and other links, plus the most granular explicit location. Do not infer location from vague context or company headquarters. **SOCIAL HANDLES:** If the CV shows a username/handle next to a social-media icon (LinkedIn, GitHub, etc.) without a full URL, reconstruct the full URL: `https://linkedin.com/in/<handle>`, `https://github.com/<handle>`, etc. Never store a bare handle like `https://username` — always include the platform domain.
7. **WORK EXPERIENCE INTEGRITY:** Merge all employment history into the single `experience` array, even if the CV splits it into multiple employment sections. Preserve explicit company, role title, dates, location, highlights, responsibilities, achievements, project details, and environment. Do not split one role unless clearly shown. Do not duplicate roles. Do not confuse role title with company name. If a role has only a duration (e.g. `6 years 5 months`) but no start/end dates, put the duration string into `dates.start` and leave `dates.end` as `""`.
8. **CURRENT TITLE & NAME NORMALIZATION:** Preserve `basics.current_title` as close as possible to the resume header wording. Extract the real display name if explicit; if not, and the email clearly contains a safe `firstname.lastname` pattern, normalize it into a human-readable name. Do not invent beyond that.
9. **CANONICAL SECTION ROUTING:** Core content must go only into its canonical sections: `basics`, `summary`, `skills`, `experience`, `education`, `certifications`, `languages`. Degrees must go to `education`, certifications to `certifications`, and language items with proficiency/test details to `languages`. **SKILLS SECTION DETECTION:** Sections titled "Technical Expertise", "Technical Skills", "Core Competencies", "Technologies", "Tech Stack", or similar MUST be parsed into the `skills` dict, NOT placed in `other_sections`. Parse "Category: item1, item2" patterns into `{"Category": ["item1", "item2"]}`.
10. **OTHER_SECTIONS ONLY:** Any remaining non-core content must go only into `other_sections`. Do not create, use, or reference `custom_sections`.
11. **CLEAN OUTPUT:** Use only empty strings `""` or arrays `[]` for missing values. Never output `None`, `null`, or placeholders. Keep wording faithful to the source but readable in professional US English. Avoid accidental ALL CAPS except for true acronyms or proper names.
12. **FIX TYPOS & GRAMMAR:** Silently fix obvious spelling mistakes, typos, grammatical errors, and incorrect product/brand capitalization. Do NOT change factual content, proper names of people, or technical terms that may look unusual but are correct.

**FINAL CHECK:**
- all human-readable text is in US English
- no explicit facts were lost
- no unsupported facts were invented
- every highlight in each role traces back to actual text in the source CV for that role
- no `None` or `null` appears
- skills from Top Skills / bullets / environment / responsibilities were not missed
- dates are preserved exactly as supported by the source; duration-only entries (e.g. `6 years 5 months`) are in `dates.start`
- no core content leaked into `other_sections`
- all remaining non-core content is preserved in `other_sections`
""",

    "prompt_qa": """Act as a strict QA Auditor for a recruitment agency. Compare the original attached CV with this extracted JSON.
NOTE: The JSON is generated programmatically. It purposefully forces skill categorization, infers standard job titles if missing, and standardizes empty fields to "". Do NOT report these architectural features as hallucinations.
NOTE: The pipeline TRANSLATES all content to US English. If the original CV is in Russian, Chinese, or any other language, the JSON will contain the English translation. This is EXPECTED behavior — do NOT flag translated content as missing or hallucinated. Compare meaning, not language.

EXTRACTED JSON:
{json_str}

TASK: Find ONLY real data losses or hallucinations. Pay special attention to education years/dates — if the original CV contains graduation years or date ranges for education entries, they MUST be present in the JSON `education[].year` field.
You MUST end your response with a JSON block in this EXACT format:
```json
{"score": 95, "missing": ["Skill 1", "Missing Date"], "hallucinations": ["Fake Certification"]}
```
If perfect, reply with score 100 and empty arrays.""",

    "prompt_autofix": """You are a Self-Healing AI. Your task is to fix the JSON extraction of a CV based on a QA Audit report.
CURRENT (BROKEN) JSON:
{current_json_str}

QA AUDIT REPORT (ERRORS TO FIX):
{qa_report_text}

INSTRUCTIONS:
1. Look at the original attached CV file/text.
2. Fix ONLY the missing data or hallucinations mentioned in the QA report.
3. Maintain the EXACT same JSON schema as the CURRENT JSON. 
4. Do NOT remove any existing correct data.
5. Do NOT add new highlights, achievements, or responsibilities that are not explicitly present in the original CV. Only restore text that was missed during extraction.
6. Place restored data into its CORRECT schema location. Education years go into `education[].year`, not into `other_sections`. Degrees go into `education[].degree`, not elsewhere.
7. Return ONLY the repaired JSON object without markdown wrappers.""",

    "prompt_matcher": """Act as a Senior IT Recruiter. Evaluate the candidate against the Job Description.
CRITICAL: Carefully analyze their actual 'experience' (duration, context, tools used), not just the 'skills' list.
First, think step-by-step about their fit.
Then, return a JSON ARRAY containing EXACTLY ONE object.

JD: {jd_val}
CANDIDATE: {cand_data}

SCHEMA:
[{
  "id": {i},
  "name": "Candidate Name",
  "reasoning": "Brief explanation of your evaluation based on their real experience",
  "score": integer (0-100),
  "verdict": "Short summary of fit",
  "pros": "Key strengths",
  "missing_skills": "Gaps or weaknesses"
}]""",

    "prompt_modifier": """You are an Expert CV Editor. Your task is to modify the provided Candidate JSON based EXACTLY on the user's request.

USER REQUEST:
"{user_req}"

CRITICAL RULES:
1. You MUST return ONLY a valid JSON object. No markdown formatting blocks, no explanations, no chat.
2. The returned JSON MUST strictly adhere to the exact same schema as the input JSON. Do not remove mandatory keys.
3. SAFE DELETION: If the user asks to remove, delete, or hide certain data (e.g., emails, phone numbers, specific jobs), DO NOT remove the keys from the JSON. Instead, set their values to an empty string `""` or an empty array `[]`.
4. AGGRESSIVE SHORTENING: If the user asks to shorten or reduce the CV (e.g., to 1 page), you MUST aggressively summarize: keep only the most recent/relevant jobs, limit achievements to 2-3 bullet points per job, and clear out secondary courses/hobbies by setting their values to empty strings/arrays.
5. Do not invent new work experience or skills unless the user asks you to infer or summarize existing ones.
6. LANGUAGE STRICTNESS: ALWAYS output the modified JSON content in professional US English.

INPUT JSON:
{input_json_str}""",

    "prompt_github": """You are an Expert Tech Recruiter. Convert this candidate's GitHub data into our STRICT JSON CV SCHEMA.
RULES:
1. Map GitHub 'name' (or 'login' if name is null) to `basics.name`.
2. Map 'location', 'email', 'blog', 'html_url' into the `basics` section.
3. Infer a professional `basics.current_title` based on their top languages.
4. Convert 'recent_repos' into the `projects` array.
5. Extract all programming languages used into the `skills` object.
6. Write a professional 2-3 sentence `summary.bullet_points` assessing their code footprint based on the "Code Quality Over Vanity" principle.
7. Leave `experience` and `education` as empty arrays `[]`.

JSON SCHEMA:
{prompt_schema_only}

GitHub API Data: {gh_full_data}""",

    "prompt_xray": """Act as an Expert Tech Recruiter.
Based on the following request, generate 3-5 advanced Google X-Ray Boolean search queries to find candidate profiles.
Target platforms: LinkedIn, GitHub, or general web.
Return ONLY a valid JSON array of objects, without Markdown formatting.
Schema:
[
  {"platform": "LinkedIn", "description": "Broad search for mid-level", "query": "site:linkedin.com/in (\\"Python\\" OR \\"Django\\") AND \\"AWS\\""}
]
Request: {user_input}""",

    "prompt_gap_analysis": """You are an expert CV-to-Job-Description fit analyst. Analyze how well this candidate matches the job requirements.

JOB DESCRIPTION:
{jd_text}

CANDIDATE CV (extracted JSON):
{cv_json}

Return a JSON object with this EXACT structure:
{{
  "match_percentage": <integer 0-100>,
  "summary": "<2-3 sentence overall fit assessment>",
  "strengths": ["<strength 1>", "<strength 2>", ...],
  "weaknesses": ["<weakness 1>", "<weakness 2>", ...],
  "skills_table": [
    {{
      "requirement": "<JD requirement or skill>",
      "category": "Must Have" or "Nice To Have",
      "status": "Covered" or "Partial" or "Missing",
      "recommendation": "<brief recommendation for tailoring, empty string if Covered>"
    }}
  ]
}}

Rules:
- strengths and weaknesses: 3-5 items each, concise bullet points
- skills_table: extract ALL identifiable requirements from the JD, classify each
- category: "Must Have" for core requirements explicitly stated as required; "Nice To Have" for preferred/bonus/optional items
- status: "Covered" if CV clearly demonstrates the skill/experience, "Partial" if related but not exact match, "Missing" if not evidenced in CV
- match_percentage: realistic assessment weighing Must Have coverage more heavily
- recommendation: actionable suggestion for how tailoring could address the gap; empty string if fully Covered
- Return ONLY valid JSON, no markdown fences""",

    "prompt_anonymize": """Act as a CV writer. Convert company names to generic industry descriptions (e.g., 'Large FinTech Company', 'Global E-commerce Enterprise'). Return JSON: {"Original": "Description"}.\nCompanies: {companies_json}""",

    "prompt_refine": """You are an Expert CV Refinement Specialist performing a SURGICAL second pass on an already-tailored CV.

JOB DESCRIPTION:
{jd_text}

MISSING JD KEYWORDS (not found in current CV):
{missing_keywords}

RULES — follow ALL strictly:
1. Your goal is to WEAVE the missing keywords into the existing text WHERE the candidate has genuine adjacent experience. For example if the CV mentions "chromatin accessibility" and "scATACseq" is missing, add it. If there is NO adjacent experience, do NOT add the keyword.
2. Do NOT fabricate experience, skills, achievements, or projects.
3. Do NOT restructure, reorder, or remove existing content.
4. PRESERVE ALL numbers exactly (publication counts, years, team sizes, percentages).
5. PRESERVE ALL dates, company names, role titles, education entries unchanged.
6. Keep changes minimal and natural — adjust wording, add synonyms, expand abbreviations. Do NOT rewrite whole sections.
7. Return ONLY a valid JSON object with exactly TWO top-level keys:
   - "_refinement_notes": 2-3 sentences listing SPECIFIC keyword insertions you made and where.
   - "cv": the complete refined CV JSON (same schema as input).
   No markdown, no explanations — raw JSON only.

INPUT JSON:
{input_json_str}""",

    "prompt_tailor": """You are an Expert CV Tailoring Specialist. Your task is to ACTIVELY REWRITE and REORDER the provided Candidate JSON to best match the given Job Description (JD). Do NOT just copy the CV — you MUST make visible changes. Do NOT invent new experience, skills, or achievements.

JOB DESCRIPTION:
{jd_text}

MANDATORY TAILORING ACTIONS (you MUST do ALL of these):

1. **OUTPUT FORMAT:** Return ONLY a valid JSON object with exactly TWO top-level keys:
   - "_tailoring_notes": 2-3 sentences listing SPECIFIC changes you made. End with "Relevance: HIGH/MEDIUM/LOW".
   - "cv": the complete tailored CV JSON (same schema as input).
   No markdown, no explanations — raw JSON only.

2. **REWRITE SUMMARY (MANDATORY):** You MUST rewrite cv.summary.bullet_points from scratch. Create 4-6 bullet points that directly connect the candidate's real experience to JD requirements. Use JD keywords naturally. Each bullet should address a specific JD requirement using the candidate's actual background. Do NOT copy the original summary unchanged.

3. **REORDER SKILLS (MANDATORY):** Move skill categories and items most relevant to the JD to the TOP. If the JD mentions AWS and the candidate has AWS buried in a list — move it to position #1. Merge small or redundant skill categories. Do NOT create a generic catch-all category like "Tools & Technologies" that duplicates items already listed in other skill groups. NEVER create empty skill categories — if the candidate has no skills for a JD-required area, simply omit that category.

4. **REORDER & REPHRASE HIGHLIGHTS (MANDATORY):** For EACH role:
   - Put JD-relevant highlights FIRST.
   - Rephrase highlights to use JD terminology where the meaning is preserved (e.g., if JD says "microservices" and highlight says "distributed systems" — use "microservices architecture").
   - Trim highlights clearly irrelevant to the JD. Set highlights to [] ONLY if ALL are irrelevant.
   - NEVER invent new highlights.

5. **PRESERVE:** Keep chronological experience order. Do NOT change: name, contacts, dates, company names, role titles, education, certifications. Keep the EXACT same JSON schema. NEVER alter stated quantities (publication counts, years of experience, team sizes, percentages, project counts) — keep the candidate's original claimed numbers even if fewer items are explicitly listed.

6. **DEDUPLICATION:** If duplicate experience entries exist (same company + similar dates), merge into one.

7. **JD TERMINOLOGY:** Where the candidate has matching experience, use the exact terms and phrases from the JD (e.g., if JD says "harmonization of biological datasets" — use that phrase, not a generic synonym). This applies to summary, skills, and experience bullets.

8. **UNIVERSAL DEVELOPER SKILLS (MANDATORY):** If the JD lists fundamental skills that any experienced developer would possess — such as data structures & algorithms, code review, version control/Git, agile/scrum, debugging, unit testing, CI/CD — you MUST add them to the output even if the source CV does not mention them explicitly. Specifically:
   - Add a mention in at least one experience highlight (e.g., "Participated in code reviews ensuring code quality and adherence to best practices", "Applied knowledge of data structures and algorithms to optimize performance-critical components").
   - Add them to the skills section under an appropriate existing category (e.g., add "Data Structures & Algorithms" to a "Core" or "Computer Science" category, add "Code Review" to a "Methodologies" or "Practices" category).
   - This is NOT fabrication — these are baseline competencies for any professional developer. Omitting them makes the CV look like a poor match when it is not.

9. Output all cv content in professional US English.

INPUT JSON:
{input_json_str}"""
}


CURRENT_PROMPT_MASTER_VERSION = 2

DEFAULT_CONFIG = {
    "api_key": "", "github_token": "", "workspace_path": DEFAULT_WORKSPACE,
    "import_mode": "qa", "anon_cut_name": True, "anon_remove_creds": True,
    "anon_mask_companies": True, "keep_initial_current_title": False,
    "show_col_file": True, "show_col_company": True, "show_col_comments": True, "show_col_score": True,
    "show_xray_tab": False, "show_github_tab": False, "show_matcher_tab": False,
    "show_modify_tab": False, "show_qa_tab": False,
    "active_template": "quantori_classic.docx", 
    "json_naming_template": "CV_FirstName_LastName.json",
    "export_naming_template": "CV_FirstName_LastName.docx",
    "naming_template": "CV FirstName FirstLetter (CV_Alexei_L.docx)",
    "ui_theme": "Light", "last_jd": "", "last_sourcing_query": "", "last_modifier_query": "",
    "generate_docx_on_import": True,
    "prompt_master_version": CURRENT_PROMPT_MASTER_VERSION, "active_prompt_version": CURRENT_PROMPT_MASTER_VERSION,
    "prompt_master_user_edited": False, "_prompt_master_upgrade_warning": False,
    "last_qa_sample_size": "All available", "qa_compare_mode": "full_pipeline", "last_miner_keywords": "", "last_miner_location": "", "last_miner_stars": "100",
    "total_in_tokens": 0, "total_out_tokens": 0, "total_spent_usd": 0.0
}
# Merge default prompts into the main config
DEFAULT_CONFIG.update(DEFAULT_PROMPTS)

def _initial_master_prompts_registry():
    return {
        "active_version": CURRENT_PROMPT_MASTER_VERSION,
        "versions": {
            str(CURRENT_PROMPT_MASTER_VERSION): {
                "version": CURRENT_PROMPT_MASTER_VERSION,
                "title": "Current production baseline",
                "status": "active",
                "based_on": None,
                "notes": "Bootstrapped from code defaults.",
                "prompt_text": DEFAULT_PROMPTS["prompt_master_inst"],
            }
        }
    }


def load_master_prompts_registry():
    if os.path.exists(MASTER_PROMPTS_FILE):
        try:
            with open(MASTER_PROMPTS_FILE, 'r', encoding='utf-8') as f:
                reg = json.load(f)
                if isinstance(reg, dict):
                    return reg
        except Exception:
            pass
    return _initial_master_prompts_registry()


def save_master_prompts_registry(registry):
    with open(MASTER_PROMPTS_FILE, 'w', encoding='utf-8') as f:
        json.dump(registry, f, indent=2, ensure_ascii=False)


def ensure_master_prompts_registry():
    reg = load_master_prompts_registry()
    if not isinstance(reg.get("versions"), dict):
        reg["versions"] = {}
    key = str(CURRENT_PROMPT_MASTER_VERSION)
    if key not in reg["versions"]:
        reg["versions"][key] = {
            "version": CURRENT_PROMPT_MASTER_VERSION,
            "title": "Current production baseline",
            "status": "active" if reg.get("active_version", CURRENT_PROMPT_MASTER_VERSION) == CURRENT_PROMPT_MASTER_VERSION else "baseline",
            "based_on": None,
            "notes": "Bootstrapped from code defaults.",
            "prompt_text": DEFAULT_PROMPTS["prompt_master_inst"],
        }
    if not reg.get("active_version"):
        reg["active_version"] = CURRENT_PROMPT_MASTER_VERSION
    save_master_prompts_registry(reg)
    return reg


def get_master_prompt_entry(version=None, registry=None):
    reg = registry or ensure_master_prompts_registry()
    if version is None:
        version = reg.get("active_version", CURRENT_PROMPT_MASTER_VERSION)
    return (reg.get("versions") or {}).get(str(version))


def get_master_prompt_text(version=None, registry=None):
    entry = get_master_prompt_entry(version, registry=registry)
    if entry and entry.get("prompt_text"):
        return entry.get("prompt_text")
    return DEFAULT_PROMPTS["prompt_master_inst"]


def get_master_prompt_versions(registry=None):
    reg = registry or ensure_master_prompts_registry()
    return sorted([int(k) for k in (reg.get("versions") or {}).keys()], reverse=True)


def set_active_master_prompt_version(version, registry=None):
    reg = registry or ensure_master_prompts_registry()
    version = int(version)
    key = str(version)
    if key not in reg.get("versions", {}):
        raise KeyError(f"Prompt version v{version} not found")
    reg["active_version"] = version
    for k, v in reg.get("versions", {}).items():
        if isinstance(v, dict):
            v["status"] = "active" if int(k) == version else (v.get("status") if v.get("status") != "active" else "archived")
    save_master_prompts_registry(reg)
    return reg


def save_master_prompt_version(prompt_text, title="Prompt Editor save", notes="", based_on=None, make_active=True, status="experimental", registry=None):
    reg = registry or ensure_master_prompts_registry()
    versions = reg.setdefault("versions", {})
    existing = [int(k) for k in versions.keys()] or [CURRENT_PROMPT_MASTER_VERSION]
    new_ver = max(existing) + 1
    versions[str(new_ver)] = {
        "version": new_ver,
        "title": title,
        "status": status,
        "based_on": int(based_on) if based_on not in (None, "") else None,
        "notes": notes,
        "prompt_text": prompt_text,
    }
    if make_active:
        reg["active_version"] = new_ver
        for k, v in versions.items():
            if isinstance(v, dict):
                v["status"] = "active" if int(k) == new_ver else (v.get("status") if v.get("status") != "active" else "archived")
    save_master_prompts_registry(reg)
    return new_ver, reg


def load_config():
    ensure_master_prompts_registry()
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                cfg = copy.deepcopy(DEFAULT_CONFIG)
                cfg.update(loaded)
                return cfg
        except Exception:
            pass
    cfg = copy.deepcopy(DEFAULT_CONFIG)
    save_config(cfg)
    return cfg


def save_config(cfg):
    with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
        json.dump(cfg, f, indent=4, ensure_ascii=False)

def init_workspace_folders(base_dir):
    if not os.path.exists(base_dir): os.makedirs(base_dir)
    folders = {
        "SOURCE": os.path.join(base_dir, 'source'),
        "JSON": os.path.join(base_dir, 'jsons'),
        "OUTPUT": os.path.join(base_dir, 'docxs'),
        "BLIND": os.path.join(base_dir, 'docxs_a'),
        "MODIFIED": os.path.join(base_dir, 'docxs_modified'),
        "REPORTS": os.path.join(base_dir, 'reports'),
        "TEMPLATES": os.path.join(base_dir, 'templates')
    }
    for folder in folders.values():
        if not os.path.exists(folder): os.makedirs(folder)
        
    target_template = os.path.join(folders["TEMPLATES"], "quantori_classic.docx")
    bundled_template = get_resource_path("quantori_classic.docx")
    if os.path.exists(bundled_template):
        # Always overwrite workspace template with the bundled one (template updates must win).
        # Keep a lightweight backup of the previous template if it already exists.
        try:
            if os.path.exists(target_template):
                bak = os.path.join(folders["TEMPLATES"], "quantori_classic.prev.docx")
                try:
                    shutil.copy2(target_template, bak)
                except Exception:
                    pass
            shutil.copy2(bundled_template, target_template)
        except Exception:
            pass
    return folders

def open_folder(path):
    if not os.path.exists(path): os.makedirs(path, exist_ok=True)
    if platform.system() == "Windows": os.startfile(path)
    elif platform.system() == "Darwin": subprocess.Popen(["open", path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    else: subprocess.Popen(["xdg-open", path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

# ==========================================
# 3. JSON SANITIZATION
# ==========================================

def fix_company_name_artifacts(data: dict) -> dict:
    """
    Fix common parsing artifacts where a section label is mistakenly captured as company_name.
    Example: company_name == "Project" / "Accomplishments" / "Environment".
    Lossless approach: we do NOT delete content, only clear the incorrect company_name value.
    """
    if not isinstance(data, dict):
        return data

    BAD = {
        "project", "projects", "project:", "projects:",
        "accomplishment", "accomplishments", "accomplishments:",
        "environment", "environment:", "responsibilities", "responsibilities:",
        "role", "role:", "position", "position:", "company", "company:",
        "time period", "time period:", "timeperiod", "timeperiod:",
    }

    def _is_bad(s: str) -> bool:
        if not isinstance(s, str):
            return False
        t = s.strip().lower()
        if not t:
            return False
        # exact matches
        if t in BAD:
            return True
        # common patterns like "Project" surrounded by punctuation
        if t.rstrip(".") in BAD:
            return True
        return False

    # v1: experience[]
    exp = data.get("experience")
    if isinstance(exp, list):
        for job in exp:
            if isinstance(job, dict):
                cn = job.get("company_name")
                if isinstance(cn, str) and _is_bad(cn):
                    job["company_name"] = ""

    # v2: work_experience[]
    wexp = data.get("work_experience")
    if isinstance(wexp, list):
        for job in wexp:
            if isinstance(job, dict):
                cn = job.get("company_name")
                if isinstance(cn, str) and _is_bad(cn):
                    job["company_name"] = ""

    return data


def _collect_raw_text_for_languages(data: dict) -> str:
    try:
        raw = data.get("raw")
        if isinstance(raw, str):
            return raw
        if isinstance(raw, dict):
            parts = []
            for v in raw.values():
                if isinstance(v, str):
                    parts.append(v)
                elif isinstance(v, list):
                    parts.extend([x for x in v if isinstance(x, str)])
            return "\n".join(parts)
        if isinstance(raw, list):
            return "\n".join([x for x in raw if isinstance(x, str)])
    except Exception:
        pass
    return ""

def _short_lang_level(s: str) -> str:
    if not isinstance(s, str):
        return ""
    t = s.strip()
    if not t:
        return ""
    # Prefer CEFR token if present
    m = re.search(r"\b([ABC][12])\b", t, flags=re.I)
    if m:
        return m.group(1).upper()
    # Normalize native variants
    if re.search(r"\bnative\b", t, flags=re.I):
        return "Native"
    return t

def ensure_native_languages(data: dict) -> dict:
    """If raw text contains 'NATIVE <LANG>' and it's not in languages[], add it losslessly."""
    if not isinstance(data, dict):
        return data
    raw_text = _collect_raw_text_for_languages(data)
    if not raw_text:
        return data
    langs = data.get("languages")
    if not isinstance(langs, list):
        return data

    existing = set()
    for it in langs:
        if isinstance(it, dict):
            name = (it.get("language") or "").strip().lower()
            if name:
                existing.add(name)
        elif isinstance(it, str):
            existing.add(it.strip().lower())

    # capture patterns like 'NATIVE RUSSIAN' or 'Native Russian'
    for m in re.finditer(r"\bNATIVE\s+([A-Z][A-Z]+|[A-Z][a-z]+)\b", raw_text, flags=re.I):
        lang = m.group(1).strip()
        if not lang:
            continue
        # Title-case if all caps
        if lang.isupper():
            lang = lang.title()
        key = lang.lower()
        if key in existing:
            continue
        if not _is_probably_tech_language(lang):
            langs.append({"language": lang, "proficiency": "Native", "level": "Native", "details": "Native"})
        existing.add(key)

    data["languages"] = langs
    return data



def _strip_leading_list_marker_text(s: str) -> str:
    """Remove leading bullet/list markers after optional indentation.

    This is intentionally conservative: it only trims marker glyphs at the
    start of a string (after whitespace). Inline dashes/bullets inside normal
    text are preserved.
    """
    if not isinstance(s, str):
        return s
    # Normalize NBSP that often appears after copied bullets.
    s = s.replace(" ", " ")
    # Repeated marker sequences like "• • AWS" / "-- item" / "— item"
    return re.sub(r'^\s*(?:[•●■▪▫◦‣⁃∙·*\-–—►▸▶➤➜❯❱→✦✧◆◇]+\s*)+', '', s)


def _strip_markdown_bold(s: str) -> str:
    """Remove markdown bold markers (**text** → text)."""
    if not isinstance(s, str):
        return s
    return re.sub(r'\*\*(.+?)\*\*', r'\1', s)


def _strip_leading_list_markers_deep(obj):
    """Recursively remove leading list markers and markdown from all text fields."""
    if isinstance(obj, str):
        return _strip_markdown_bold(_strip_leading_list_marker_text(obj))
    if isinstance(obj, list):
        return [_strip_leading_list_markers_deep(x) for x in obj]
    if isinstance(obj, dict):
        return {k: _strip_leading_list_markers_deep(v) for k, v in obj.items()}
    return obj


def _normalize_human_language(name: str) -> str:
    """Normalize language name for whitelist matching."""
    if not isinstance(name, str):
        return ""
    s = name.strip()
    if not s:
        return ""
    # remove leading bullets and punctuation
    s = re.sub(r"^[\s•●■▪▫◦‣⁃∙·–—\-►▸▶➤➜❯❱→✦✧◆◇]+", "", s).strip()
    # drop CEFR / level in parentheses or after dash, e.g. "English (C1)" / "English - C1"
    s = re.sub(r"\([^)]*\)", "", s).strip()
    s = re.sub(r"\s*[-–—]\s*(A1|A2|B1|B2|C1|C2|Native|Fluent|Advanced|Intermediate|Beginner)\b.*$", "", s, flags=re.I).strip()
    # drop trailing commas/colons
    s = s.rstrip(" ,;:")
    return s

_HUMAN_LANG_ALIASES = {
    # abbreviations
    "en": "english",
    "eng": "english",
    "ru": "russian",
    "rus": "russian",
    "de": "german",
    "ger": "german",
    "fr": "french",
    "es": "spanish",
    "spa": "spanish",
    "pt": "portuguese",
    "por": "portuguese",
    "it": "italian",
    "nl": "dutch",
    "uk": "ukrainian",
    "ua": "ukrainian",
    "hy": "armenian",
    "ka": "georgian",
    "zh": "chinese",
    "ja": "japanese",
    "jp": "japanese",
    "ko": "korean",
    "ar": "arabic",
    "fa": "persian",
    "he": "hebrew",
}

_HUMAN_LANG_WHITELIST = {
    # top common
    "english","russian","german","french","spanish","portuguese","italian","dutch","ukrainian","polish","turkish",
    "armenian","georgian","azerbaijani","kazakh","uzbek","tajik","kyrgyz","belarusian","moldovan","romanian","bulgarian",
    "greek","serbian","croatian","bosnian","slovenian","slovak","czech","hungarian","albanian","macedonian","montenegrin",
    "swedish","norwegian","danish","finnish","icelandic","estonian","latvian","lithuanian",
    "chinese","mandarin","cantonese","japanese","korean","vietnamese","thai","indonesian","malay","filipino","tagalog",
    "hindi","urdu","bengali","punjabi","tamil","telugu","marathi","gujarati","kannada","malayalam","sinhala","nepali",
    "arabic","persian","farsi","hebrew","kurdish",
    "swahili","afrikaans","zulu","xhosa","amharic",
    "latin","irish","scottish gaelic","welsh",
    # variants
    "persion","french","spanish","portuguese","brazilian portuguese","brazilian","português",
}

def _is_human_language(name: str) -> bool:
    """Strict whitelist-based classifier for human languages."""
    n0 = _normalize_human_language(name)
    if not n0:
        return False
    n = n0.lower()
    # normalize common alias tokens (two/three-letter codes)
    if n in _HUMAN_LANG_ALIASES:
        n = _HUMAN_LANG_ALIASES[n]
    # normalize 'farsi' => persian
    if n == "farsi":
        n = "persian"
    # normalize multiple spaces
    n = re.sub(r"\s+", " ", n).strip()
    return n in _HUMAN_LANG_WHITELIST

def sync_languages_to_skills(data: dict) -> dict:
    """Render human Languages into skills['Languages'] with levels, losslessly.
    Strict: only whitelist human languages go into this list.
    Non-matching tokens are preserved in extras as 'Languages(unclassified): ...' (lossless).

    Additionally, remove duplicates like 'English' + 'English (C1)' by keeping the richer version.
    """
    if not isinstance(data, dict):
        return data

    langs = data.get("languages")
    if not isinstance(langs, list) or not langs:
        return data

    def _base_name(s: str) -> str:
        s = _normalize_human_language(s)
        s = re.sub(r"\s+", " ", s).strip().lower()
        if s in _HUMAN_LANG_ALIASES:
            s = _HUMAN_LANG_ALIASES[s]
        if s == "farsi":
            s = "persian"
        return s

    rendered = []
    unclassified = []

    for it in langs:
        if isinstance(it, dict):
            name = (it.get("language") or "").strip()
            if not name:
                continue
            if not _is_human_language(name):
                unclassified.append(name)
                continue
            lvl = _short_lang_level(it.get("level") or it.get("proficiency") or it.get("details") or "")
            rendered.append(f"{name} ({lvl})" if lvl else name)
        elif isinstance(it, str) and it.strip():
            name = it.strip()
            if not _is_human_language(name):
                unclassified.append(name)
                continue
            rendered.append(name)

    # Merge previous Languages entries (only human ones), but don't re-introduce poorer duplicates.
    skills = data.get("skills")
    if not isinstance(skills, dict):
        skills = {}

    prev = skills.get("Languages")
    if isinstance(prev, list):
        for x in prev:
            if isinstance(x, str) and x.strip():
                nm = x.strip()
                if _is_human_language(nm):
                    rendered.append(nm)

    # Choose best variant per base language name, preserving first-seen order of bases.
    best_by_base = {}
    order = []
    for x in rendered:
        base = _base_name(x)
        if not base:
            continue
        if base not in best_by_base:
            best_by_base[base] = x
            order.append(base)
        else:
            cur = best_by_base[base]
            # Prefer strings that include a level "(C1)" / "(Native)" or are longer.
            def score(v: str) -> int:
                s = 0
                if re.search(r"\([ABC][12]\)", v):
                    s += 5
                if re.search(r"\(Native\)", v, flags=re.I):
                    s += 4
                if "(" in v and ")" in v:
                    s += 2
                s += min(3, len(v)//10)
                return s
            if score(x) > score(cur):
                best_by_base[base] = x

    rendered2 = [best_by_base[b] for b in order if b in best_by_base]

    skills["Languages"] = rendered2
    data["skills"] = skills

    # Lossless: stash unclassified tokens so they are not lost
    if unclassified:
        extras = data.get("extras")
        if not isinstance(extras, list):
            extras = []
        for tok in unclassified:
            line = f"Languages(unclassified): {tok}"
            if line not in extras:
                extras.append(line)
        data["extras"] = extras

    return data



_MONTH_NAMES = ["january","february","march","april","may","june",
                "july","august","september","october","november","december"]

# Titles that duplicate canonical sections — must be filtered from other_sections
# to prevent double-rendering (LinkedIn PDFs often contain these as extra sections)
CANONICAL_SECTION_TITLES = {
    "work experience", "experience", "опыт работы", "опыт",
    "education", "образование",
    "skills", "technical skills", "top skills", "навыки", "технические навыки",
    "основные навыки", "ключевые навыки", "core competencies", "key skills",
    "languages", "языки",
    "summary", "summary of qualifications", "professional summary", "career summary",
    "profile", "objective", "резюме", "о себе", "общие сведения",
    "certifications", "сертификаты", "сертификации",
    "contacts", "contact information", "contact", "способы связаться", "контакты",
    "links", "ссылки",
}

def _has_non_ascii(s):
    """Return True if string contains non-ASCII characters (likely non-English)."""
    return bool(s) and any(ord(c) > 127 for c in s)


def translate_locations_via_llm(data, api_key):
    """Translate all non-English locations in CV data to English via a single LLM call.
    Modifies data in-place. Returns list of translations made (for logging)."""
    if not api_key or not isinstance(data, dict):
        return []

    locs = {}
    bl = data.get('basics', {}).get('location', '')
    if isinstance(bl, str) and _has_non_ascii(bl):
        locs.setdefault(bl.strip(), []).append((data['basics'], 'location'))
    contacts = data.get('basics', {}).get('contacts', {})
    if isinstance(contacts, dict):
        cl = contacts.get('location', '')
        if isinstance(cl, str) and _has_non_ascii(cl):
            locs.setdefault(cl.strip(), []).append((contacts, 'location'))
    for job in data.get('experience', []):
        jl = job.get('location', '')
        if isinstance(jl, str) and _has_non_ascii(jl):
            locs.setdefault(jl.strip(), []).append((job, 'location'))
    for edu in data.get('education', []):
        for k in ('location', 'details'):
            el = edu.get(k, '')
            if isinstance(el, str) and _has_non_ascii(el):
                locs.setdefault(el.strip(), []).append((edu, k))

    if not locs:
        return []

    unique_locs = list(locs.keys())
    numbered = "\n".join(f"{i+1}. {loc}" for i, loc in enumerate(unique_locs))
    prompt = (
        "Translate the following location names to US English. "
        "Return ONLY a JSON array of translated strings in the same order. "
        "Keep the original comma-separated structure (city, country). "
        "If already in English, return as-is.\n\n"
        f"{numbered}"
    )

    try:
        client = genai.Client(api_key=api_key)
        resp = client.models.generate_content(model=MODEL_NAME, contents=prompt)
        txt = (getattr(resp, 'text', '') or '').strip().replace('```json', '').replace('```', '').strip()
        translated = json.loads(txt)
        if not isinstance(translated, list) or len(translated) != len(unique_locs):
            return []
        changes = []
        for orig, eng in zip(unique_locs, translated):
            eng = str(eng).strip()
            if eng and eng != orig:
                for obj, key in locs[orig]:
                    obj[key] = eng
                changes.append(f"{orig} → {eng}")
        return changes
    except Exception:
        return []


def translate_dates_via_llm(data, api_key):
    """Translate all non-English dates in CV data to English via a single LLM call.
    Modifies data in-place. Returns list of translations made (for logging)."""
    if not api_key or not isinstance(data, dict):
        return []

    dates_map = {}
    for section_key in ('experience', 'education', 'certifications'):
        items = data.get(section_key)
        if not isinstance(items, list):
            continue
        for item in items:
            if not isinstance(item, dict):
                continue
            dates = item.get('dates')
            if isinstance(dates, dict):
                for k in ('start', 'end'):
                    v = dates.get(k, '')
                    if isinstance(v, str) and _has_non_ascii(v):
                        dates_map.setdefault(v.strip(), []).append((dates, k))
            for k in ('date', 'period', 'year'):
                v = item.get(k, '')
                if isinstance(v, str) and _has_non_ascii(v):
                    dates_map.setdefault(v.strip(), []).append((item, k))

    if not dates_map:
        return []

    unique_dates = list(dates_map.keys())
    numbered = "\n".join(f"{i+1}. {d}" for i, d in enumerate(unique_dates))
    prompt = (
        "Translate the following date strings to US English. "
        "Use format like 'January 2023', '03/2020', 'Present' etc. "
        "If a date means 'current/now/present', translate as 'Present'. "
        "Return ONLY a JSON array of translated strings in the same order. "
        "If already in English, return as-is.\n\n"
        f"{numbered}"
    )

    try:
        client = genai.Client(api_key=api_key)
        resp = client.models.generate_content(model=MODEL_NAME, contents=prompt)
        txt = (getattr(resp, 'text', '') or '').strip().replace('```json', '').replace('```', '').strip()
        translated = json.loads(txt)
        if not isinstance(translated, list) or len(translated) != len(unique_dates):
            return []
        changes = []
        for orig, eng in zip(unique_dates, translated):
            eng = str(eng).strip()
            if eng and eng != orig:
                for obj, key in dates_map[orig]:
                    obj[key] = eng
                changes.append(f"{orig} → {eng}")
        return changes
    except Exception:
        return []


def _has_cyrillic(s):
    """Return True if string contains Cyrillic characters."""
    return bool(s) and any('\u0400' <= c <= '\u04FF' for c in str(s))


def _count_cyrillic_strings(obj):
    """Count how many string values in a nested structure contain Cyrillic."""
    count = 0
    if isinstance(obj, str):
        return 1 if _has_cyrillic(obj) else 0
    if isinstance(obj, list):
        for item in obj:
            count += _count_cyrillic_strings(item)
    elif isinstance(obj, dict):
        for k, v in obj.items():
            if k.startswith('_') or k in ('qa_audit', 'import_date', '_source_filename', '_source_hash'):
                continue
            count += _count_cyrillic_strings(v)
    return count


def translate_full_json_via_llm(data, api_key):
    """Translate entire CV JSON to US English if significant Cyrillic content detected.
    Also removes duplicate entries (Russian+English pairs in experience/education/certifications).
    Modifies data in-place. Returns description of changes or empty string."""
    if not api_key or not isinstance(data, dict):
        return ""

    cyrillic_count = _count_cyrillic_strings(data)
    if cyrillic_count < 3:
        return ""  # Minor Cyrillic (e.g. just a location) — handled by specific translators

    # Prepare a clean copy without volatile keys
    SKIP_KEYS = {'qa_audit', 'match_analysis', '_status', 'selected', 'ts',
                 'import_date', '_source_filename', '_source_hash', '_comment'}
    clean = {k: v for k, v in data.items() if k not in SKIP_KEYS}

    prompt = (
        "You are a professional CV translator. The JSON below contains a CV with mixed "
        "Russian and English content, and possibly DUPLICATE entries (same role/education/certification "
        "appearing in both Russian and English).\n\n"
        "Your tasks:\n"
        "1. TRANSLATE all Russian text to professional US English.\n"
        "2. DEDUPLICATE: if the same experience role, education entry, or certification appears "
        "in both Russian and English, keep ONLY the English version (merging any extra details from the Russian one).\n"
        "3. PRESERVE all factual content — do NOT remove, invent, or alter any facts.\n"
        "4. Keep the exact same JSON structure and keys.\n"
        "5. Return ONLY the translated JSON object, no markdown wrappers.\n\n"
        f"{json.dumps(clean, indent=2, ensure_ascii=False)}"
    )

    try:
        client = genai.Client(api_key=api_key)
        resp = client.models.generate_content(model=MODEL_NAME, contents=prompt)
        txt = (getattr(resp, 'text', '') or '').strip()
        txt = txt.replace('```json', '').replace('```', '').strip()
        translated = json.loads(txt)

        if not isinstance(translated, dict):
            return ""

        # Verify no significant data loss: translated should not have fewer strings
        # (but may have fewer due to dedup — allow some reduction)
        orig_strs = _count_non_empty_strings(clean)
        new_strs = _count_non_empty_strings(translated)
        # Allow up to 40% fewer strings (dedup removes Russian duplicates)
        if new_strs < orig_strs * 0.5:
            return ""  # Too much lost — reject

        # Check Cyrillic is actually reduced
        new_cyrillic = _count_cyrillic_strings(translated)
        if new_cyrillic >= cyrillic_count:
            return ""  # Translation didn't help

        # Apply translated content back to data (preserve skipped keys)
        for k in list(data.keys()):
            if k not in SKIP_KEYS and k in translated:
                data[k] = translated[k]

        return f"Translated {cyrillic_count} non-English fields to US English (Cyrillic remaining: {new_cyrillic})"
    except Exception:
        return ""


def _count_non_empty_strings(obj):
    """Count non-empty strings in a nested structure."""
    count = 0
    if isinstance(obj, str):
        return 1 if obj.strip() else 0
    if isinstance(obj, list):
        for item in obj:
            count += _count_non_empty_strings(item)
    elif isinstance(obj, dict):
        for v in obj.values():
            count += _count_non_empty_strings(v)
    return count


def translate_remaining_strings_via_llm(data, api_key):
    """Find and translate any remaining non-English string values in CV JSON.
    Handles stray Cyrillic/CJK/Arabic/etc. strings that weren't caught by
    full translation or date/location translators.
    Modifies data in-place. Returns list of changes for logging."""
    if not api_key or not isinstance(data, dict):
        return []

    SKIP_KEYS = {'qa_audit', 'match_analysis', '_status', 'selected', 'ts',
                 'import_date', '_source_filename', '_source_hash', '_comment'}

    # Collect all non-English strings with their paths for in-place update
    non_eng = {}  # text -> list of (parent_obj, key_or_index)

    def _walk(obj, parent=None, key=None):
        if isinstance(obj, str):
            if obj.strip() and _has_non_ascii(obj):
                non_eng.setdefault(obj.strip(), []).append((parent, key))
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                _walk(item, obj, i)
        elif isinstance(obj, dict):
            for k, v in obj.items():
                if k in SKIP_KEYS or k.startswith('_'):
                    continue
                _walk(v, obj, k)

    _walk(data)

    if not non_eng:
        return []

    unique_strings = list(non_eng.keys())
    numbered = "\n".join(f"{i+1}. {s}" for i, s in enumerate(unique_strings))

    prompt = (
        "Translate the following strings to professional US English. "
        "Keep technical terms, proper names, and abbreviations as-is. "
        "Return ONLY a JSON array of translated strings in the same order. "
        "If a string is already in English, return it unchanged.\n\n"
        f"{numbered}"
    )

    try:
        client = genai.Client(api_key=api_key)
        resp = client.models.generate_content(model=MODEL_NAME, contents=prompt)
        txt = (getattr(resp, 'text', '') or '').strip()
        txt = txt.replace('```json', '').replace('```', '').strip()
        translated = json.loads(txt)

        if not isinstance(translated, list) or len(translated) != len(unique_strings):
            return []

        changes = []
        for orig, eng in zip(unique_strings, translated):
            eng = str(eng).strip()
            if eng and eng != orig:
                for parent, key in non_eng[orig]:
                    if parent is not None and key is not None:
                        parent[key] = eng
                changes.append(f"{orig[:40]} → {eng[:40]}")
        return changes
    except Exception:
        return []


def _is_future_date(s):
    """Return True if date string represents a month/year clearly in the future."""
    if not isinstance(s, str): return False
    low = s.strip().lower()
    if not low or low == 'present': return False
    m = re.search(r'\b(20\d\d|19\d\d)\b', s)
    if not m: return False
    year = int(m.group(1))
    today = datetime.date.today()
    if year > today.year: return True
    if year == today.year:
        for i, name in enumerate(_MONTH_NAMES):
            if name in low:
                return (i + 1) > today.month
    return False


def sanitize_json(data):
    if not isinstance(data, dict): data = {}
    data = _strip_leading_list_markers_deep(data)

    backup_keys = {k: data.get(k) for k in ['qa_audit', 'match_analysis', '_source_filename', '_source_hash', 'import_date', '_comment']}

    if not isinstance(data.get('basics'), dict): data['basics'] = {}
    bad_values = ['unavailable', 'n/a', 'none', 'null', 'not provided', 'unknown', 'undisclosed']
    
    for key in ['name', 'location', 'objective', 'current_company']:
            val = data['basics'].get(key)
            if isinstance(val, str):
                cleaned_val = val.strip().replace('\n', ' ').replace('\r', '')
                if key == 'name' and cleaned_val:
                    if cleaned_val.isupper() or cleaned_val.islower(): cleaned_val = cleaned_val.title()
                data['basics'][key] = "" if cleaned_val.lower() in bad_values else cleaned_val
            else: 
                data['basics'][key] = ", ".join(map(str, val)) if isinstance(val, list) else (str(val) if val else "")

    raw_title = data['basics'].get('current_title', '')
    if isinstance(raw_title, str):
            raw_title = raw_title.replace('\n', ' ').replace('\r', '').strip()
            if raw_title.lower() in bad_values: raw_title = ""
    else: raw_title = str(raw_title) if raw_title else ""
            
    data['basics']['current_title_original'] = raw_title
            
    _title_role_kw = {'engineer', 'manager', 'developer', 'analyst', 'designer', 'architect',
                      'lead', 'director', 'consultant', 'specialist', 'scientist', 'officer',
                      'coordinator', 'executive', 'head', 'vp', 'president', 'intern', 'sdet',
                      'administrator', 'researcher', 'devops', 'founder', 'partner', 'cto', 'ceo'}
    clean_title = raw_title
    if clean_title:
            # Strip common prefixes that LLM sometimes prepends
            clean_title = re.sub(r'^(Objective|Summary|Profile|About)\s*:\s*', '', clean_title, flags=re.IGNORECASE).strip()
            # Extract and rescue tech terms from parenthesised content before stripping
            _paren_match = re.search(r'\((.+)', clean_title)
            if _paren_match:
                _paren_text = re.sub(r'[()]+', ',', _paren_match.group(1))
                _paren_text = re.sub(r'\b\d+x\s+', '', _paren_text)
                _paren_text = re.sub(r'\b(expert|certified|specialist|proficient|experienced)\b', '', _paren_text, flags=re.IGNORECASE)
                _paren_text = re.sub(r',?\s*[BM]\.\s?[SA]\.?\s+(in\s+\w+)+.*$', '', _paren_text, flags=re.IGNORECASE)
                _title_techs = [t.strip() for t in re.split(r'[,&]+', _paren_text) if t.strip() and len(t.strip()) >= 2]
                if _title_techs:
                    skills = data.setdefault('skills', {})
                    existing = set()
                    for cat_items in skills.values():
                        if isinstance(cat_items, list):
                            existing.update(s.lower().strip() for s in cat_items)
                    new_techs = [t for t in _title_techs if t.lower().strip() not in existing]
                    if new_techs:
                        title_cat = skills.get('Title Specialties', [])
                        title_cat.extend(new_techs)
                        skills['Title Specialties'] = title_cat
            # Remove parenthesised tech lists / credential noise before splitting
            clean_title = re.sub(r'\s*\(.*', '', clean_title).strip()
            # Remove trailing degree/credential fragments like "M.S in ComputerScience"
            clean_title = re.sub(r',?\s*[BM]\.\s?[SA]\.?\s+(in\s+\w+)+.*$', '', clean_title, flags=re.IGNORECASE).strip()
            segments = re.split(r'\s*\|\s*|\s*-\s*|\s*,\s*|\s+at\s+|\s+@\s+', clean_title, flags=re.IGNORECASE)
            # Pick the first segment that looks like a job title; fall back to first segment
            clean_title = segments[0].strip()
            for seg in segments:
                seg = seg.strip()
                if seg and any(kw in seg.lower() for kw in _title_role_kw):
                    clean_title = seg
                    break
            if clean_title.isupper(): clean_title = clean_title.title()
            if len(clean_title) > 80 or len(clean_title) < 2: clean_title = ""
            # Reject sentence-like text (objective/summary leaked into title)
            if clean_title and re.search(r'\b(years?\s+of|experience\s+in|responsible\s+for|worked\s+(in|at|on)|passionate\s+about)\b', clean_title, re.IGNORECASE):
                clean_title = ""

    # If title looks like a location (e.g. "Bay Area/San Diego") rather than a job title, discard it
    if clean_title and '/' in clean_title:
        if not any(kw in clean_title.lower() for kw in _title_role_kw):
            clean_title = ""

    if not clean_title and data.get('experience'):
            role_raw = data['experience'][0].get('role')
            fb = str(role_raw).replace('\n', ' ').strip() if role_raw else ''
            if fb.lower() in bad_values: fb = ''
            if fb.isupper(): fb = ' '.join(w if (w.isupper() and len(w) <= 4) else w.capitalize() for w in fb.split())
            if fb: clean_title = fb
    # Last resort: use objective if it looks like a job title (short, no sentences)
    if not clean_title:
            obj = data['basics'].get('objective', '')
            if isinstance(obj, str) and 0 < len(obj) <= 80 and not re.search(r'\b(years?\s+of|experience\s+in)\b', obj, re.IGNORECASE):
                clean_title = obj

    data['basics']['current_title'] = clean_title

    if not isinstance(data['basics'].get('contacts'), dict): data['basics']['contacts'] = {}
    clean_contacts = {}
    for k, v in data['basics']['contacts'].items():
        s_val = ", ".join(map(str, v)) if isinstance(v, list) else str(v)
        if s_val and s_val.lower().strip() not in bad_values: clean_contacts[k] = s_val
    data['basics']['contacts'] = clean_contacts
    
    if not isinstance(data['basics'].get('links'), list): data['basics']['links'] = []
    raw_links = [str(l).strip() for l in data['basics']['links'] if l and str(l).lower().strip() not in bad_values]
    # Fix links: expand bare handles, fix missing-domain URLs, deduplicate
    clean_links = []
    seen_links = set()
    _social_platform_hints = {
        'linkedin': 'https://linkedin.com/in/',
        'in': 'https://linkedin.com/in/',
        'github': 'https://github.com/',
        'gh': 'https://github.com/',
    }
    for lnk in raw_links:
        expanded = []
        # Normalize: strip leading "@", "in", "linkedin", "github" prefixes to extract handle
        # e.g. "in @magnitopic" → handle "magnitopic", hint "linkedin"
        # e.g. "@magnitopic" → handle "magnitopic", no hint
        # e.g. "https://magnitopic" → handle "magnitopic"
        cleaned = lnk
        detected_platform = None
        # Strip known platform prefix words
        for prefix, url_base in _social_platform_hints.items():
            if re.match(rf'^{prefix}\b\s*', cleaned, re.IGNORECASE):
                cleaned = re.sub(rf'^{prefix}\s*', '', cleaned, flags=re.IGNORECASE).strip()
                detected_platform = url_base
                break
        # Strip @ and protocol
        cleaned = re.sub(r'^@', '', cleaned).strip()
        cleaned = re.sub(r'^https?://', '', cleaned).strip()
        # Now check if it's a bare handle (alphanumeric, no dots/spaces)
        if re.match(r'^[a-zA-Z0-9_-]+$', cleaned):
            if detected_platform:
                expanded = [f"{detected_platform}{cleaned}"]
                # Also add the other platform
                other = 'https://github.com/' if 'linkedin' in detected_platform else 'https://linkedin.com/in/'
                expanded.append(f"{other}{cleaned}")
            else:
                expanded = [f"https://linkedin.com/in/{cleaned}", f"https://github.com/{cleaned}"]
        # Full URL with domain (contains a dot)
        elif '.' in lnk:
            expanded = [lnk]
        # Unrecognized format — keep as-is
        else:
            expanded = [lnk]
        for url in expanded:
            low = url.lower()
            if low not in seen_links:
                seen_links.add(low)
                clean_links.append(url)
    data['basics']['links'] = clean_links

    if 'skills' not in data or not isinstance(data['skills'], dict): data['skills'] = {}
    clean_skills = {}
    for k, v in data['skills'].items():
        # Normalize schema-leaked key names: "technical_skills" -> "Technical Skills"
        clean_k = k.replace('_', ' ').strip().title() if '_' in k else k
        if isinstance(v, str): clean_skills[clean_k] = [v]
        elif isinstance(v, list): clean_skills[clean_k] = [str(i) for i in v if i]
    data['skills'] = clean_skills

    if not isinstance(data.get('experience'), list): data['experience'] = []
    for job in data['experience']:
        pd = job.get('project_description')
        if isinstance(pd, list): pd = " ".join(map(str, pd))
        if not isinstance(pd, str): pd = ""
        pd_stripped = pd.strip()
        # Discard values that are pure punctuation/whitespace (e.g. ":", "-", "—")
        if pd_stripped.lower() in bad_values or (pd_stripped and not re.search(r'\w', pd_stripped)):
            pd_stripped = ""
        job['project_description'] = pd_stripped
            
        loc = job.get('location')
        job['location'] = "" if (isinstance(loc, str) and loc.lower().strip() in bad_values) else (loc if isinstance(loc, str) else "")
        
        hl = job.get('highlights')
        job['highlights'] = [hl] if isinstance(hl, str) else (hl if isinstance(hl, list) else [])
        
        env = job.get('environment')
        job['environment'] = [env] if isinstance(env, str) else (env if isinstance(env, list) else [])
        
        if 'dates' in job:
            for d_key in ['start', 'end']:
                val = job['dates'].get(d_key)
                job['dates'][d_key] = "" if (isinstance(val, str) and val.lower().strip() in bad_values) else (val if isinstance(val, str) else "")
            # Future end date → replace with "Present" (LLM hallucination guard)
            end_val = job['dates'].get('end', '')
            if _is_future_date(end_val):
                job['dates']['end'] = 'Present'
        
        c_val = job.get('company_name')
        job['company_name'] = "" if (isinstance(c_val, str) and c_val.lower().strip() in bad_values) else (c_val if isinstance(c_val, str) else "")

        # Normalize ALLCAPS roles (e.g. "SENIOR SOFTWARE ENGINEER" → "Senior Software Engineer")
        role_val = job.get('role', '')
        if isinstance(role_val, str) and role_val == role_val.upper() and len(role_val) > 3:
            _role_acronyms = {'QA', 'VP', 'CTO', 'CIO', 'CFO', 'CEO', 'IT', 'HR', 'PM', 'DBA',
                              'SDET', 'UX', 'UI', 'DevOps', 'SRE', 'ML', 'AI', 'BA', 'BI'}
            words = re.split(r'(\s+|/)', role_val)  # preserve separators
            job['role'] = ''.join(w if w.strip() in _role_acronyms else w.title() for w in words)

    # Deduplicate experience entries:
    # 1) Exact match by (company_name, role, start_date)
    # 2) Fuzzy match: same/similar dates AND overlapping company name
    # 3) Same company + same role (catches LLM duplicating with different date formats)
    def _norm_date(d):
        """Normalize date string for comparison: strip spaces/dashes, extract year digits."""
        return re.sub(r'[^a-z0-9]', '', str(d).strip().lower())

    seen_exp = set()
    clean_exp = []
    for job in data['experience']:
        key = (
            str(job.get('company_name', '')).strip().lower(),
            str(job.get('role', '')).strip().lower(),
            str((job.get('dates') or {}).get('start', '')).strip().lower(),
        )
        if key in seen_exp:
            continue
        seen_exp.add(key)
        # Fuzzy duplicate check
        j_start = _norm_date((job.get('dates') or {}).get('start', ''))
        j_end = _norm_date((job.get('dates') or {}).get('end', ''))
        j_company = str(job.get('company_name', '')).strip().lower()
        j_role = str(job.get('role', '')).strip().lower()
        is_fuzzy_dup = False
        if j_company:
            for existing in clean_exp:
                e_start = _norm_date((existing.get('dates') or {}).get('start', ''))
                e_end = _norm_date((existing.get('dates') or {}).get('end', ''))
                e_company = str(existing.get('company_name', '')).strip().lower()
                e_role = str(existing.get('role', '')).strip().lower()
                # Match criteria: overlapping company name AND either:
                #   a) normalized dates match, OR
                #   b) same role text (catches date-shuffled duplicates)
                company_match = e_company and (j_company in e_company or e_company in j_company or j_company == e_company)
                dates_match = j_start and e_start and (j_start == e_start or j_start in e_start or e_start in j_start) and (j_end == e_end or j_end in e_end or e_end in j_end)
                role_match = j_role and e_role and j_role == e_role
                if company_match and (dates_match or role_match):
                    # Merge: keep the entry with more highlights; add missing highlights from duplicate
                    e_hl = existing.get('highlights') or []
                    j_hl = job.get('highlights') or []
                    if len(j_hl) > len(e_hl):
                        existing['highlights'] = j_hl
                    elif j_hl:
                        e_hl_lower = {h.lower().strip() for h in e_hl}
                        for h in j_hl:
                            if h.lower().strip() not in e_hl_lower:
                                e_hl.append(h)
                    # Keep the longer company name (more specific)
                    if len(j_company) > len(e_company):
                        existing['company_name'] = job.get('company_name', '')
                    # Keep non-empty project_description
                    if not existing.get('project_description') and job.get('project_description'):
                        existing['project_description'] = job['project_description']
                    # Merge environments
                    e_env = existing.get('environment') or []
                    j_env = job.get('environment') or []
                    if j_env:
                        e_env_lower = {x.lower().strip() for x in e_env}
                        for x in j_env:
                            if x.lower().strip() not in e_env_lower:
                                e_env.append(x)
                        existing['environment'] = e_env
                    is_fuzzy_dup = True
                    break
        if not is_fuzzy_dup:
            clean_exp.append(job)
    data['experience'] = clean_exp

    # Enrich skills from experience environments: collect unique items not already in skills
    existing_skills_lower = set()
    for items in data.get('skills', {}).values():
        if isinstance(items, list):
            existing_skills_lower.update(s.lower().strip() for s in items if isinstance(s, str))
    env_skills = []
    for job in data.get('experience', []):
        for item in job.get('environment', []):
            if isinstance(item, str) and item.strip():
                if item.lower().strip() not in existing_skills_lower:
                    env_skills.append(item.strip())
                    existing_skills_lower.add(item.lower().strip())
    if env_skills:
        # Append to existing "Tools & Technologies" or create it
        target_key = None
        for k in data['skills']:
            if k.lower() in ('tools & technologies', 'tools and technologies', 'technologies', 'tools'):
                target_key = k
                break
        if target_key:
            data['skills'][target_key].extend(env_skills)
        else:
            data['skills']['Tools & Technologies'] = env_skills

    # Enrich experience environment from highlights: extract tech terms mentioned inline
    # Only match well-defined tech terms (tools, frameworks, languages) using word boundaries
    _known_tech_lower = {}  # lowercase -> original case
    for items in data.get('skills', {}).values():
        if isinstance(items, list):
            for s in items:
                if isinstance(s, str) and len(s.strip()) > 2:
                    _known_tech_lower[s.lower().strip()] = s.strip()
    for job in data.get('experience', []):
        env = job.get('environment') or []
        env_lower = {e.lower().strip() for e in env}
        text_parts = list(job.get('highlights') or [])
        if job.get('project_description'):
            text_parts.append(job['project_description'])
        combined_text = ' '.join(str(t) for t in text_parts)
        for tech_lower, tech_orig in _known_tech_lower.items():
            if tech_lower not in env_lower:
                # Use word boundary matching to avoid partial matches
                if re.search(r'\b' + re.escape(tech_lower) + r'\b', combined_text, re.IGNORECASE):
                    env.append(tech_orig)
                    env_lower.add(tech_lower)
        job['environment'] = env

    if not isinstance(data.get('projects'), list): data['projects'] = []
    clean_projects = []
    for p in data['projects']:
        if not isinstance(p, dict): p = {}
        for k in ['title', 'description', 'link']:
            val = p.get(k)
            p[k] = "" if (isinstance(val, str) and val.lower().strip() in bad_values) else (val if isinstance(val, str) else "")
            
        if 'tech_stack' not in p or not isinstance(p['tech_stack'], list): p['tech_stack'] = []
        p['tech_stack'] = [str(t) for t in p['tech_stack'] if t and str(t).lower().strip() not in bad_values]
        if p.get('title') or p.get('description'): clean_projects.append(p)
    data['projects'] = clean_projects
        
    for k in ['certifications', 'publications', 'courses']:
        if k not in data or not isinstance(data[k], list): data[k] = []
        else: data[k] = [str(x) for x in data[k] if x and str(x).lower().strip() not in bad_values]

    if 'education' in data and isinstance(data['education'], list):
        for edu in data['education']:
            for e_key in ['institution', 'degree', 'year', 'details']:
                val = edu.get(e_key)
                if val is None:
                    edu[e_key] = ""
                elif isinstance(val, str) and val.lower().strip() in bad_values:
                    edu[e_key] = ""
        # Deduplicate education entries by (institution, degree, year)
        seen_edu = set()
        clean_edu = []
        for edu in data['education']:
            key = (
                str(edu.get('institution', '')).strip().lower(),
                str(edu.get('degree', '')).strip().lower(),
                str(edu.get('year', '')).strip().lower(),
            )
            if key in seen_edu:
                continue
            seen_edu.add(key)
            clean_edu.append(edu)
        data['education'] = clean_edu

    if 'volunteering' not in data or not isinstance(data['volunteering'], list): data['volunteering'] = []
    for v in data['volunteering']:
        if not isinstance(v, dict): continue
        for k in ['organization', 'role']:
            val = v.get(k)
            v[k] = "" if (isinstance(val, str) and val.lower().strip() in bad_values) else (str(val) if val else "")
        if 'highlights' not in v or not isinstance(v['highlights'], list): v['highlights'] = []
        else: v['highlights'] = [str(h) for h in v['highlights'] if h]

    data = normalize_languages_lossless(data)
    if 'languages' not in data or not isinstance(data['languages'], list): data['languages'] = []
    for l in data['languages']:
        if not isinstance(l, dict): continue
        # keep extra fields (level/details), but sanitize placeholders
        for k in ['language','proficiency','level','details']:
            if k not in l: continue
            val = l.get(k)
            l[k] = "" if (isinstance(val, str) and val.lower().strip() in bad_values) else (str(val) if val else "")

    # Merge any legacy custom_sections into other_sections, then remove the legacy field
    # so other_sections remains the only non-core bucket in the final JSON.
    if not isinstance(data.get('other_sections'), list):
        data['other_sections'] = []

    def _is_mojibake(s):
        """Return True if string looks like UTF-8 bytes misread as Latin-1 (>40% in U+0080-U+00FF)."""
        if not s:
            return False
        extended = sum(1 for c in s if '\x80' <= c <= '\xff')
        return extended / len(s) > 0.4

    def _normalize_other_section(sec, *, title_keys=("title", "section_title")):
        if not isinstance(sec, dict):
            return None
        title = ""
        for key in title_keys:
            raw_title = sec.get(key, "")
            if isinstance(raw_title, str):
                raw_title = raw_title.strip()
            elif raw_title:
                raw_title = str(raw_title).strip()
            else:
                raw_title = ""
            if raw_title and _is_mojibake(raw_title):
                raw_title = ""
            if raw_title:
                title = _normalize_optional_section_title(raw_title)
                break
        if isinstance(title, str) and title.lower().strip() in bad_values:
            title = ""

        raw_items = sec.get('items', [])
        if isinstance(raw_items, list):
            items = [str(x).strip() for x in raw_items if x and str(x).strip() and str(x).lower().strip() not in bad_values]
        elif raw_items and str(raw_items).strip() and str(raw_items).lower().strip() not in bad_values:
            items = [str(raw_items).strip()]
        else:
            items = []

        if not title and not items:
            return None
        return {"title": title, "items": items}

    # Rescue skills-like sections from other_sections into skills dict
    _skills_like_titles = {'technical expertise', 'technical skills', 'core competencies',
                           'technologies', 'tech stack', 'tools and technologies',
                           'tools & technologies', 'competencies', 'key skills',
                           'professional skills', 'areas of expertise'}
    rescued_other = []
    for sec in data.get('other_sections', []):
        norm = _normalize_other_section(sec, title_keys=("title", "section_title"))
        if not norm:
            continue
        sec_title_lower = norm["title"].strip().lower()
        if sec_title_lower in _skills_like_titles and norm.get("items"):
            # Parse "Category: item1, item2" patterns into skills dict
            for item in norm["items"]:
                if ':' in item:
                    cat, vals = item.split(':', 1)
                    cat = cat.strip()
                    parsed = [v.strip() for v in vals.split(',') if v.strip()]
                    if parsed:
                        if cat in data['skills']:
                            existing_lower = {s.lower() for s in data['skills'][cat]}
                            for v in parsed:
                                if v.lower() not in existing_lower:
                                    data['skills'][cat].append(v)
                        else:
                            data['skills'][cat] = parsed
                else:
                    # Single item without category — add to general
                    data['skills'].setdefault('General', []).append(item.strip())
            continue
        rescued_other.append(norm)
    data['other_sections'] = rescued_other

    # If basics.summary is empty, rescue content from summary-like other_sections before filtering
    _summary_like = {'summary', 'summary of qualifications', 'professional summary',
                     'career summary', 'executive summary', 'profile summary'}
    merged_other = []
    for sec in data.get('other_sections', []):
        norm = _normalize_other_section(sec, title_keys=("title", "section_title"))
        if not norm:
            continue
        sec_title_lower = norm["title"].strip().lower()
        # Rescue summary content before discarding canonical sections
        if not data['basics'].get('summary') and sec_title_lower in _summary_like and norm.get("items"):
            data['basics']['summary'] = "\n".join(norm["items"])
            continue
        if sec_title_lower not in CANONICAL_SECTION_TITLES:
            merged_other.append(norm)
    # Migrate legacy top-level non-core fields into other_sections
    # so other_sections becomes the only canonical non-core bucket.
    def _project_lines(projects):
        lines = []
        if not isinstance(projects, list):
            return lines
        for p in projects:
            if isinstance(p, dict):
                title = str(p.get("title") or p.get("name") or "").strip()
                desc = str(p.get("description") or "").strip()
                link = str(p.get("link") or p.get("url") or "").strip()
                line = title or ""
                if link:
                    line = (line + " — " + link).strip(" —")
                if desc:
                    line = (line + " — " + desc).strip(" —")
                if line:
                    lines.append(line)
            elif isinstance(p, str) and p.strip():
                lines.append(p.strip())
        return lines

    legacy_map = [
        ("projects", "Projects", _project_lines),
        ("courses", "Courses", lambda v: [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []),
        ("publications", "Publications", lambda v: [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []),
        ("volunteering", "Volunteering", lambda v: [
            " — ".join([
                str(it.get("organization", "")).strip(),
                str(it.get("role", "")).strip()
            ]).strip(" —")
            for it in v if isinstance(it, dict) and (
                str(it.get("organization", "")).strip() or str(it.get("role", "")).strip()
            )
        ] if isinstance(v, list) else []),
        ("extras", "Other", lambda v: [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []),
        ("other", "Other", lambda v: [str(x).strip() for x in v if str(x).strip()] if isinstance(v, list) else []),
    ]

    for key, title, fn in legacy_map:
        lines = fn(data.get(key))
        if lines:
            merged_other.append({"title": title, "items": lines})
    legacy_custom = data.get('custom_sections', [])
    if isinstance(legacy_custom, list):
        for sec in legacy_custom:
            norm = _normalize_other_section(sec, title_keys=("section_title", "title"))
            if norm:
                merged_other.append(norm)

    # De-duplicate by (title, items) while preserving order.
    seen_sections = set()
    clean_other = []
    for sec in merged_other:
        sig = (sec.get('title', '').strip().casefold(), tuple(i.casefold() for i in sec.get('items', [])))
        if sig in seen_sections:
            continue
        seen_sections.add(sig)
        clean_other.append(sec)

    # Remove duplicated core sections from other_sections if canonical core is already filled
    filtered_other = []
    for sec in clean_other:
        title = (sec.get("title", "") or "").strip().casefold()
        if title == "languages" and data.get("languages"):
            continue
        if title == "certifications" and data.get("certifications"):
            continue
        if title == "education" and data.get("education"):
            continue
        filtered_other.append(sec)

    # Rescue orphaned education years from other_sections into education[].year
    # Autofix sometimes puts years like "(2012 - 2016)" into other_sections instead of education
    _year_re = re.compile(r'\(?\s*(\d{4})\s*[-–—]\s*(\d{4})\s*\)?$')
    edu_list = data.get('education', [])
    if edu_list:
        rescued_items = set()
        for sec in filtered_other:
            for idx, item in enumerate(sec.get('items', [])):
                m = _year_re.match(item.strip())
                if not m:
                    continue
                year_str = f"{m.group(1)} - {m.group(2)}"
                # Find an education entry with empty year
                for edu in edu_list:
                    if not edu.get('year'):
                        edu['year'] = year_str
                        rescued_items.add((id(sec), idx))
                        break
        # Remove rescued items from other_sections
        if rescued_items:
            new_filtered = []
            for sec in filtered_other:
                new_items = [item for idx, item in enumerate(sec.get('items', []))
                             if (id(sec), idx) not in rescued_items]
                if new_items or sec.get('title'):
                    sec['items'] = new_items
                    if new_items:
                        new_filtered.append(sec)
                else:
                    pass  # drop empty section
            filtered_other = new_filtered

    data['other_sections'] = filtered_other

    # Remove legacy non-core containers from final JSON
    for legacy_key in [
        'custom_sections',
        'projects',
        'courses',
        'publications',
        'volunteering',
        'extras',
        'other',
    ]:
        data.pop(legacy_key, None)

    # Remove obvious section-label artifacts (e.g., company_name == 'Project')
    data = fix_company_name_artifacts(data)

    for k, v in backup_keys.items():
        if v is not None: data[k] = v

    return data



def normalize_languages_lossless(data: dict) -> dict:
    """
    Lossless normalization for languages:
    - Accept list items as dict or string
    - Preserve extra fields (level/details/notes)
    - Extract CEFR-like level from proficiency/details when possible
    """
    if not isinstance(data, dict):
        return data
    langs = data.get("languages")
    if langs is None:
        return data
    if not isinstance(langs, list):
        langs = [langs]
    out = []
    # CEFR / common level tokens
    level_re = re.compile(r"\b(A1|A2|B1|B2|C1|C2)\b", re.I)
    native_re = re.compile(r"\b(native|mother\s*tongue|fluent|bilingual)\b", re.I)

    for item in langs:
        if isinstance(item, str):
            s = item.strip()
            if not s:
                continue
            # Try split "Language - Level/Details"
            # Examples: "English — C1 (Advanced)", "Russian (Native)"
            lang = s
            prof = ""
            # split on dash-like separators
            parts = re.split(r"\s*(?:[-–—:|]|,)\s*", s, maxsplit=1)
            if len(parts) == 2 and parts[0] and parts[1]:
                lang = parts[0].strip()
                prof = parts[1].strip()
            d = {"language": lang, "proficiency": prof}
            out.append(d)
            continue

        if isinstance(item, dict):
            # keep everything, but ensure canonical keys exist
            d = dict(item)
            lang = str(d.get("language") or d.get("name") or "").strip()
            prof = str(d.get("proficiency") or d.get("level") or "").strip()
            details = str(d.get("details") or d.get("notes") or "").strip()

            # If language is missing but we have a single key dict like {"English": "C1"} treat it
            if not lang and len(d) == 1:
                k = next(iter(d.keys()))
                v = d[k]
                if isinstance(k, str):
                    lang = k.strip()
                    prof = str(v).strip() if v is not None else prof

            # Extract level from prof/details
            level = str(d.get("level") or "").strip()
            if not level:
                m = level_re.search(prof) or level_re.search(details)
                if m:
                    level = m.group(1).upper()
                else:
                    m2 = native_re.search(prof) or native_re.search(details)
                    if m2:
                        level = m2.group(1).title()

            # If details empty but prof contains more than just level, keep it
            if not details and prof and (len(prof) > 4 or "(" in prof or "/" in prof):
                details = prof

            d["language"] = lang
            d["proficiency"] = prof
            if level:
                d["level"] = level
            if details:
                d["details"] = details

            # Drop known bad placeholders but keep lossless extras
            if lang or prof or details or level:
                out.append(d)
            continue

    data["languages"] = out
    
    # Languages: add missing Native entries from raw and render levels into skills['Languages']
    data = ensure_native_languages(data)
    data = sync_languages_to_skills(data)

    return data

def _collect_raw_text(ctx: dict) -> str:
    parts = []
    raw = ctx.get("raw") or {}
    if isinstance(raw, dict):
        for k in ("source_text_chunks","unmapped_facts"):
            v = raw.get(k)
            if isinstance(v, list):
                parts.extend([x for x in v if isinstance(x,str)])
            elif isinstance(v, str):
                parts.append(v)
    for k in ("source_text","raw_text","text","_source_text"):
        v = ctx.get(k)
        if isinstance(v, str):
            parts.append(v)
    return "\n".join([p.strip() for p in parts if isinstance(p,str) and p.strip()])

def extract_text_from_docx(docx_path: str) -> str:
    """Robust DOCX text extraction for CVs (paragraphs + tables + headers/footers)."""
    try:
        from docx import Document
    except Exception:
        return ""
    try:
        doc = Document(docx_path)
    except Exception:
        return ""

    parts = []

    def add_paras(paras):
        for p in paras:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

    add_paras(doc.paragraphs)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                add_paras(cell.paragraphs)

    try:
        for sec in doc.sections:
            add_paras(sec.header.paragraphs)
            add_paras(sec.footer.paragraphs)
    except Exception:
        pass

    # dedupe consecutive identical lines
    cleaned = []
    prev = None
    for line in parts:
        if line == prev:
            continue
        cleaned.append(line)
        prev = line

    return "\n".join(cleaned).strip()

# ==========================================
# 4. LLM PROCESSING CORE
# ==========================================
def _format_docx_sections_for_llm(baseline: dict) -> str:
    """Format structured DOCX sections into labeled text for the LLM."""
    sections = baseline.get("sections", {})
    section_labels = {
        "preamble": "[CONTACT HEADER: extract name/email/phone/location into basics fields ONLY — do NOT use as current_title, do NOT create a section]",
        "summary": "SUMMARY",
        "skills": "TECHNICAL SKILLS",
        "experience": "WORK EXPERIENCE",
        "education": "EDUCATION",
        "certifications": "CERTIFICATIONS",
        "languages": "LANGUAGES",
    }
    parts = []
    for key, label in section_labels.items():
        lines = sections.get(key, [])
        if lines:
            parts.append(f"=== {label} ===\n" + "\n".join(lines))
    for key, lines in sections.items():
        if key not in section_labels and lines:
            parts.append(f"=== {key.upper().replace('_', ' ')} ===\n" + "\n".join(lines))
    return "\n\n".join(parts)


def process_file_gemini(file_path, api_key, custom_instructions, task_state=None):
    # 🧠 Combine editable instructions and the protected schema
    final_prompt = custom_instructions + f"\n\n**JSON SCHEMA:**\n{CV_JSON_SCHEMA}"

    if file_path.lower().endswith('.docx'):
        try:
            from source_baseline_extractor import extract_from_docx as _extract_from_docx
            baseline = _extract_from_docx(file_path)
            text = _format_docx_sections_for_llm(baseline)
        except Exception:
            text = extract_text_from_docx(file_path)
        if not text: raise ValueError("Empty DOCX")
        response = _generate_text_content_with_retry(api_key, [final_prompt, text])
    else:
        mime = 'application/pdf' if file_path.lower().endswith('.pdf') else 'image/jpeg'
        sample = _upload_gemini_file_and_wait(api_key, file_path, mime, task_state=task_state)
        if sample is None:
            return None, 0, 0, 0.0
        if task_state and task_state.get("cancel"): return None, 0, 0, 0.0
        response = _generate_file_content_with_retry(api_key, final_prompt, sample)
        
    text = response.text.replace('```json', '').replace('```', '').strip()
    
    in_tok, out_tok = _extract_token_usage(response)
    cost = (in_tok / 1_000_000 * PRICE_1M_IN) + (out_tok / 1_000_000 * PRICE_1M_OUT)
    
    if not text: return None, in_tok, out_tok, cost
    from converter_engine import extract_first_json_object
    data = extract_first_json_object(text)
    return sanitize_json(data), in_tok, out_tok, cost

def generate_docx_from_json(data, output_path, cfg):
    template_name = cfg.get("active_template", "quantori_classic.docx")

    # Priority 1: absolute path already resolved by converter_engine
    if cfg.get("template_path") and os.path.exists(cfg["template_path"]):
        template_path = cfg["template_path"]
    else:
        # Priority 2: workspace templates (desktop / local dev)
        workspace = cfg.get("workspace_path", DEFAULT_WORKSPACE)
        candidate = os.path.join(workspace, "templates", template_name)
        # Priority 3: templates/ next to cv_engine.py (Docker / server deployments)
        engine_dir = os.path.dirname(os.path.abspath(__file__))
        local_candidate = os.path.join(engine_dir, "templates", template_name)
        local_classic = os.path.join(engine_dir, "templates", "quantori_classic.docx")

        if os.path.exists(candidate):
            template_path = candidate
        elif os.path.exists(local_candidate):
            template_path = local_candidate
        elif os.path.exists(local_classic):
            template_path = local_classic
        else:
            raise FileNotFoundError(f"Template not found! Please create: {candidate}")
        
    doc = DocxTemplate(template_path)
    context = _trim_strings_deep(copy.deepcopy(data))
    # Ensure required keys exist for Jinja templates (avoid StrictUndefined crashes)
    if not isinstance(context, dict):
        context = {}
    context.setdefault('basics', {})
    context.setdefault('summary', {})
    context.setdefault('skills', {})
    context.setdefault('experience', [])
    context.setdefault('education', [])
    context.setdefault('certifications', [])
    context.setdefault('projects', [])
    context.setdefault('extras', [])
    
    c_list = [str(v).strip() for k, v in context.get('basics', {}).get('contacts', {}).items() if str(v).strip()]
    if str(context.get('basics', {}).get('location', '')).strip(): 
        c_list.append(str(context['basics']['location']).strip())
    c_list.extend([str(x).strip() for x in context.get('basics', {}).get('links', []) if str(x).strip()])
    context['contact_line'] = " | ".join(c_list)
    
    if cfg.get('keep_initial_current_title', False):
        context['basics']['current_title'] = context.get('basics', {}).get('current_title_original', context.get('basics', {}).get('current_title', ''))
    # Optional: enrich contact_line from raw text if empty
    raw_text = _collect_raw_text(context)
    if not context.get('contact_line'):
        c2 = _extract_contacts_plus(raw_text)
        parts = []
        for k in ('email', 'phone', 'website', 'linkedin'):
            v = c2.get(k)
            if isinstance(v, str) and v.strip():
                parts.append(v.strip())
        loc = _extract_location_line(raw_text)
        if isinstance(loc, str) and loc.strip():
            parts.append(loc.strip())
        if parts:
            # de-dupe while preserving order
            context['contact_line'] = " | ".join(dict.fromkeys(parts))

    
    # Summary normalization: move bullet points into Summary (regulations have no "Key Highlights")
    context.setdefault("summary", {})
    if not isinstance(context["summary"], dict):
        context["summary"] = {}

    # v2 -> v1 mapping
    items = context["summary"].get("items") or []
    if items and not context["summary"].get("bullet_points"):
        context["summary"]["bullet_points"] = [x.strip() for x in items if isinstance(x, str) and x.strip()]

    # If still no bullets, derive from objective (best-effort)
    if not context["summary"].get("bullet_points"):
        obj = (context.get("basics") or {}).get("objective")
        if isinstance(obj, str) and obj.strip():
            parts = re.split(r"(?<=[.!?])\s+", obj.strip())
            bullets = [p.strip() for p in parts if p.strip()]
            # limit to 7
            context["summary"]["bullet_points"] = bullets[:7]
            # Only when objective is the sole source of content — clear it afterwards to avoid
            # the template rendering both objective and summary with identical text.
            context["basics"]["objective"] = ""


    # Remove "Languages" key from skills dict to prevent double rendering
    # (once under Technical Skills, once in the dedicated Languages section).
    skills = context.get("skills")
    if isinstance(skills, dict):
        keys_to_remove = [k for k in skills if k.strip().lower() == "languages"]
        for k in keys_to_remove:
            del skills[k]

    # Build other_sections for any non-required content (goes to the end of CV under its own headings)
    # Required sections per regulations are rendered earlier (Summary/Technical Skills/Work Experience/Education/Certifications).
    other_sections = []

#=========================================================
    # Candidate optional sections (canonical only)
    # Languages: render ONLY explicit human/spoken languages from the canonical
    # structured `languages` section. Never derive this section from skills["Languages"].
    def _as_lines(lang_items):
        out = []
        if not isinstance(lang_items, list):
            return out
        for it in lang_items:
            if isinstance(it, dict):
                lang = str(it.get("language") or "").strip()
                prof = str(it.get("level") or it.get("proficiency") or it.get("details") or "").strip()
                if lang and prof:
                    out.append(f"{lang} ({prof})")
                elif lang:
                    out.append(lang)
            elif isinstance(it, str) and it.strip():
                out.append(it.strip())
        return out

    lang_lines = _as_lines(context.get("languages"))
    if lang_lines:
        other_sections.append({"title": "Languages", "items": lang_lines})

#===========================================================================            

    # Merge any pre-normalized other_sections from JSON/context so template uses a single
    # bucket for all non-core sections.
    existing_other = context.get("other_sections")
    if isinstance(existing_other, list):
        for sec in existing_other:
            if not isinstance(sec, dict):
                continue
            title = str(sec.get("title", "")).strip()
            # Skip sections that duplicate canonical sections (e.g. LinkedIn "Top Skills",
            # "Technical Skills" raw blocks, Russian section headers from older imports)
            if title.lower() in CANONICAL_SECTION_TITLES:
                continue
            title = _normalize_optional_section_title(sec.get("title", ""))
            items = sec.get("items", [])
            if not isinstance(items, list):
                items = [items] if items else []
            lines = [str(x).strip() for x in items if x and str(x).strip()]
            if title or lines:
                other_sections.append({"title": title, "items": lines})

    # Normalize ALLCAPS titles to Title Case (word-level check for mixed case like "SUMMARY of QUALIFICATIONS")
    for sec in other_sections:
        t = sec.get("title", "")
        if t and len(t) > 3:
            words = t.split()
            upper_count = sum(1 for w in words if w == w.upper() and len(w) > 1)
            if upper_count > len(words) / 2:
                sec["title"] = t.title()

    deduped_other = []
    seen_other = set()
    for sec in other_sections:
        sig = (sec.get("title", "").strip().casefold(), tuple(i.casefold() for i in sec.get("items", [])))
        if sig in seen_other:
            continue
        seen_other.add(sig)
        deduped_other.append(sec)

    context["other_sections"] = deduped_other
    doc.render(context)
    try:
        doc.save(output_path)
        return output_path
    except PermissionError:
        base, ext = os.path.splitext(output_path)
        for counter in range(1, DOCX_SAVE_MAX_RETRIES):
            target_path = f"{base}_{counter:02d}{ext}"
            try:
                doc.save(target_path)
                return target_path
            except PermissionError:
                continue
        raise PermissionError(f"Cannot save DOCX after {DOCX_SAVE_MAX_RETRIES} attempts: {output_path}")





def smart_anonymize_data(data, api_key, cfg):
    blind = copy.deepcopy(data)
    in_tok, out_tok, cost = 0, 0, 0.0
    
    if cfg.get("anon_cut_name", True):
        name = blind.get('basics', {}).get('name', '')
        if name:
            parts = name.split()
            # Preserve academic degrees (PhD, Ph.D., MD, etc.) in anonymized name
            _degree_patterns = {'phd', 'ph.d.', 'ph.d', 'md', 'm.d.', 'dsc', 'd.sc.', 'dr.'}
            degrees = [p for p in parts if p.lower().rstrip('.,') in _degree_patterns]
            name_parts = [p for p in parts if p.lower().rstrip('.,') not in _degree_patterns]
            anon_name = f"{name_parts[0]} {name_parts[1][0]}." if len(name_parts) > 1 else "Candidate"
            # If no degree in name, check education for PhD/MD
            if not degrees:
                for edu in blind.get('education', []):
                    deg = str(edu.get('degree', '')).lower()
                    if 'phd' in deg or 'ph.d' in deg or 'doctorate' in deg or 'доктор' in deg:
                        degrees.append('PhD')
                        break
                    elif deg.startswith('md') or 'm.d.' in deg:
                        degrees.append('MD')
                        break
            if degrees:
                # Deduplicate
                seen = set()
                unique_degrees = []
                for d in degrees:
                    dl = d.lower().rstrip('.,')
                    if dl not in seen:
                        seen.add(dl)
                        unique_degrees.append(d)
                anon_name = f"{anon_name}, {' '.join(unique_degrees)}"
            blind['basics']['name'] = anon_name

    if cfg.get("anon_remove_creds", True):
        if 'contacts' in blind.get('basics', {}): blind['basics']['contacts'] = {}
        if 'links' in blind.get('basics', {}): blind['basics']['links'] = []

    # Anonymize publications: replace list with summary count
    for section_list in [blind.get('other_sections', [])]:
        for sec in section_list:
            if not isinstance(sec, dict):
                continue
            title_lower = str(sec.get('title', '')).strip().lower()
            if any(kw in title_lower for kw in ('publication', 'paper', 'conference proceeding')):
                pub_count = len(sec.get('items', []))
                if pub_count > 0:
                    sec['items'] = [f"Author and co-author of {pub_count} publications in peer-reviewed scientific journals and conference proceedings."]

    if cfg.get("anon_mask_companies", True):
        experiences = blind.get('experience', [])
        companies = [job.get('company_name') for job in experiences if job.get('company_name')]
        volunteering = blind.get('volunteering', [])
        v_orgs = [v.get('organization') for v in volunteering if v.get('organization')]
        all_companies = companies + v_orgs
        
        if all_companies:
            unique_comps = list(set(all_companies))
            
            # Use the editable prompt from the config
            prompt_template = cfg.get("prompt_anonymize", DEFAULT_PROMPTS["prompt_anonymize"])
            prompt = prompt_template.replace("{companies_json}", json.dumps(unique_comps, ensure_ascii=False))
            
            try:
                response = _generate_text_content_with_retry(api_key, prompt)
                
                i_tok, o_tok = _extract_token_usage(response)
                in_tok += i_tok
                out_tok += o_tok
                cost += (i_tok / 1_000_000 * PRICE_1M_IN) + (o_tok / 1_000_000 * PRICE_1M_OUT)
                
                text = response.text.replace('```json', '').replace('```', '').strip()
                mapping = json.loads(text)
            except Exception:
                mapping = {name: "Confidential Company" for name in unique_comps}

            # Also collect client/project names mentioned in brackets in project descriptions
            bracket_names = set()
            for job in experiences:
                pd = str(job.get('project_description', ''))
                for hl in (job.get('highlights') or []):
                    pd += ' ' + str(hl)
                # Match [Name] or [Name (details)]
                for m in re.findall(r'\[([^\]]+)\]', pd):
                    name = m.strip()
                    if name and len(name) > 2 and name not in mapping:
                        bracket_names.add(name)
            # Add bracket names to mapping as "Confidential Client"
            for bn in bracket_names:
                if bn not in mapping:
                    mapping[bn] = "Confidential Client"

            # Replace company names in company_name field
            for job in experiences:
                original = job.get('company_name')
                if original and original in mapping: job['company_name'] = mapping[original]
                elif original: job['company_name'] = "Confidential Company"

            for vol in volunteering:
                original = vol.get('organization')
                if original and original in mapping: vol['organization'] = mapping[original]
                elif original: vol['organization'] = "Confidential Organization"

            # Also scrub company names from text fields (project_description, highlights, environment)
            def _scrub_text(text, name_mapping):
                if not isinstance(text, str):
                    return text
                for orig, repl in name_mapping.items():
                    if orig and orig in text:
                        text = text.replace(orig, repl)
                return text

            def _scrub_deep(obj, name_mapping):
                if isinstance(obj, str):
                    return _scrub_text(obj, name_mapping)
                if isinstance(obj, list):
                    return [_scrub_deep(item, name_mapping) for item in obj]
                if isinstance(obj, dict):
                    return {k: _scrub_deep(v, name_mapping) for k, v in obj.items()}
                return obj

            for job in experiences:
                for field in ('project_description', 'highlights', 'environment'):
                    if field in job:
                        job[field] = _scrub_deep(job[field], mapping)

    return blind, in_tok, out_tok, cost
